import json
import re
import time
from io import BytesIO
import requests
import uuid

from docx import Document


def init_client():
    client = requests.session()
    raw_headers = """Authorization: eyJ0eXAiOiJKV1QiLCJhbGciOiJIUzI1NiJ9.eyJ1c2VyX2lkIjoiNzY1ODE5NiIsInBob25lIjoiMTg5Njg5OTU4ODEiLCJjcmVhdGVfdGltZSI6MTYyMDQ0MTEzNSwidXBkYXRlX3RpbWUiOjE2MjA0NDExMzUsInN5c3RlbSI6IjI5IiwicGxhdGZvcm0iOiJBbmRyb2lkIiwidmVyc2lvbiI6IjQuMTIuMCIsInNka192ZXJzaW9uIjoiNC4xMi4wIiwibmlja25hbWUiOiIxODk2ODk5NTg4MSIsImF2YXRhciI6IiIsImdlbmRlciI6ImYiLCJzc29faWQiOjIwNjMyOTkyLCJ3ZWlib19pZCI6IiIsInFxX2lkIjoiIiwid2VpeGluX2lkIjoiIiwid2VpYm9fbmlja25hbWUiOiIiLCJxcV9uaWNrbmFtZSI6IiIsIndlaXhpbl9uaWNrbmFtZSI6IiIsImxvZ2luX3R5cGUiOiJwaG9uZSIsInVzZXJfZnJvbSI6MTUsImNvZGUiOiJHQ1RTVCIsImlhdCI6MTYyMDQ0MTEzNSwibmJmIjoxNjIwNDQxMTM1LCJleHAiOjE2NTE5NzcxMzV9.TQp3Y4KD-o_X_8TL-LNeKCyQ0ZIMsIXWsrAZLETDm4A
    Referer: http://tiku.eoffcn.com/apiv3/
    Content-Type: application/x-www-form-urlencoded
    Accept-Encoding: gzip
    User-Agent: okhttp/4.2.2"""
    for s in raw_headers.split("\n"):
        key, val = list(map(str.strip, s.split(": ")))
        client.headers[key] = val

    return client


class Paper(object):
    MockPaperReportURL = 'http://tiku.eoffcn.com/apiv3/mock/exercise/getMockPaperReport'

    def __init__(self, client, process, record_sub_id, mock_subject_id, user_id='7658196'):
        self._record_sub_id = record_sub_id
        self._mock_subject_id = mock_subject_id
        self._user_id = user_id
        self._client = client
        self._process = process

    def run(self):
        payload = {
            'record_sub_id': self._record_sub_id,
            'mock_subject_id': self._mock_subject_id,
            'channel': 15,
            'version': '4.12.0',
            'appid': 'zgjiaoyu',
            'platform': 'Android',
            'format': 'form',
            'sign': '14af3b7ccb71cd5f5c2c8b552aa106c8'
        }

        r = self._client.post(self.MockPaperReportURL, data=payload)
        if r.status_code != 200:
            print("获取该模拟考的所有题目数据错误！")
            return

        # 处理返回结果
        raw_all_problem_ids = json.loads(r.text)
        print(raw_all_problem_ids['data']['mock_title'])  # 第几期等信息
        ids = []
        for raw_list in raw_all_problem_ids['data']['list']:
            for raw in raw_list["list"]:
                ids.append(raw['question_id'])


class Process(object):
    QuestionDetail = 'http://tiku.eoffcn.com/apiv3/exampaper/wrongquestion/getWrongQuestion'

    def __init__(self, client, record_id, user_id='7658196'):
        self._client = client
        self._record_id = record_id
        self._user_id = user_id
        self._name = None

    def get_detail(self, ids):
        data = {'userId': self._user_id,
                'recordId': self._record_id,
                'origin': '3',
                'questionIds': '856437,1072831,1072832,1072833,1072834,1072835,1072836',
                'examId': '71',
                'channel': '15',
                'submit_time': '0',
                'product': '3',
                'appid': 'zgjiaoyu',
                'version': '4.12.0',
                'platform': 'Android',
                'format': 'form',
                'sign': '4a951b04ee2b83aa43d13eb4ab085afd'}

        err_cnt = 0
        for i in range(0, len(ids), 7):
            data['questionIds'] = ','.join(map(str, ids[i:i + 7]))
            r = self._client.post(self.QuestionDetail, data=data)

            if r.status_code != 200 and err_cnt > 3:
                err_cnt += 1
                print("获取题目的解析错误大于3次，请检查 %s" % ids[i:i + 7])
                return False
            yield r.text

        return True

    def _correct(self, choices):
        for i, c in enumerate(choices):
            if c['is_correct']:
                return 'ABCD'[i]
        print("错误选项")
        return 'E'

    def download(self, url):
        r = self._client.get(url)
        print(url[20:], r.status_code)
        return BytesIO(r.content)

    def subject(self, subjects):
        return ' '.join(s['first_name'] for s in subjects)

    def process(self, ids, name):
        ans = []
        document = Document()
        problem_number = 1
        for text in self.get_detail(ids):
            data = json.loads(text)
            for ph in data["data"]:
                p1 = document.add_paragraph()
                p1.add_run(
                    '%s     分值：%s    归类：%s\n' % (problem_number, ph['score'], self.subject(ph['subject']))
                ).bold = True
                ans.append(self._correct(ph['choices']))
                text = '正确答案：%s\n' % (self._correct(ph['choices']))
                p1.add_run(text)

                explanation = ph['explanation'].replace('<br>', '\n')
                if '<img' in explanation:
                    urls = re.findall('<img src="(.*?)".*?>', explanation)
                    texts = re.split('<img.*?>', explanation)
                    for t, u in zip(texts[1:], urls):
                        r = p1.add_run()
                        inline_shape = r.add_picture(self.download(u))
                        inline_shape.height = int(inline_shape.height * 0.4)
                        inline_shape.width = int(inline_shape.width * 0.5)
                        p1.add_run(t)
                        time.sleep(0.5)
                else:
                    p1.add_run('%s\n' % explanation)
                p1.style = 'Normal'
                problem_number += 1
            time.sleep(2)

        i = 1
        p1 = document.add_paragraph()
        for index in range(0, len(ans), 5):
            p1.add_run('%s-%s %s    ' % (i, i + 4, ''.join(ans[index:index + 5]))).bold = True
            i += 5

        if name:
            document.save(name + '.docx')
        else:
            document.save(str(uuid.uuid4().hex) + '.docx')
        return True

def read_config():
    with open("./config.json", 'r', encoding='UTF-8') as fp:
        return json.load(fp)

if __name__ == '__main__':
    client = init_client()
    paper = Paper(client, record_sub_id, mock_subject_id)
    process = Process(client, record_sub_id)
