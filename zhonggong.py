import json
import re
import time
from io import BytesIO
import requests
import uuid

from docx import Document


def init_client(auth):
    _client = requests.session()
    raw_headers = """Referer: http://tiku.eoffcn.com/apiv3/
    Content-Type: application/x-www-form-urlencoded
    Accept-Encoding: gzip
    User-Agent: okhttp/4.2.2"""
    for s in raw_headers.split("\n"):
        key, val = list(map(str.strip, s.split(": ")))
        _client.headers[key] = val

    _client.headers['Authorization'] = auth

    return _client


class Paper(object):
    MockPaperReportURL = 'http://tiku.eoffcn.com/apiv3/mock/exercise/getMockPaperReport'

    def __init__(self, client, process, user_id='7658196'):
        self._user_id = user_id
        self._client = client
        self._process = process

    def run(self, papers):
        for data in papers:
            record_sub_id, mock_subject_id = data['record_sub_id'], data['mock_subject_id']
            payload = {
                'record_sub_id': int(record_sub_id),
                'mock_subject_id': int(mock_subject_id),
                'channel': 15,
                'version': '4.12.0',
                'appid': 'zgjiaoyu',
                'platform': 'Android',
                'format': 'form',
                'sign': '64018318e691e99cff36437891383268'
            }

            r = self._client.post(self.MockPaperReportURL, data=payload)
            if r.status_code != 200:
                print("获取该模拟考的所有题目数据错误！")
                return False

            # 处理返回结果
            raw_all_problem_ids = json.loads(r.text)
            mock_title = raw_all_problem_ids['data']['mock_title']  # 第几期等信息
            print(mock_title)
            ids = []
            for raw_list in raw_all_problem_ids['data']['list']:
                for raw in raw_list["list"]:
                    ids.append(raw['question_id'])

            self._process(ids, record_sub_id, mock_title)
            print(mock_title, 'done')


class Process(object):
    QuestionDetail = 'http://tiku.eoffcn.com/apiv3/exampaper/wrongquestion/getWrongQuestion'

    def __init__(self, client, user_id='7658196'):
        self._client = client
        self._user_id = user_id
        self._name = None

    def get_detail(self, ids, record_id):
        data = {'userId': self._user_id,
                'recordId': record_id,
                'origin': '3',
                'questionIds': '856437,1072831,1072832,1072833,1072834,1072835,1072836',
                'examId': '71',
                'channel': '15',
                'submit_time': '0',
                'product': '3',
                'appid': 'zgjiaoyu',
                'version': '4.12.0',
                'platform': 'Android',
                'format': 'form',
                'sign': '64018318e691e99cff36437891383268'}

        err_cnt = 0
        for i in range(0, len(ids), 7):
            data['questionIds'] = ','.join(map(str, ids[i:i + 7]))
            r = self._client.post(self.QuestionDetail, data=data)

            if r.status_code != 200 and err_cnt > 3:
                err_cnt += 1
                print("获取题目的解析错误大于3次，请检查 %s" % ids[i:i + 7])
                return False
            yield r.text

        return True

    def _correct(self, choices):
        for i, c in enumerate(choices):
            if c['is_correct']:
                return 'ABCD'[i]
        print("错误选项")
        return 'E'

    def download(self, url):
        r = self._client.get(url)
        print(url[30:], r.status_code)
        return BytesIO(r.content)

    def subject(self, subjects):
        return ' '.join(s['first_name'] for s in subjects)

    def process(self, ids, record_id, name):
        ans = []
        document = Document()
        problem_number = 1
        for text in self.get_detail(ids, record_id):
            data = json.loads(text)
            for ph in data["data"]:
                p1 = document.add_paragraph()
                p1.add_run(
                    '%s     分值：%s    归类：%s\n' % (problem_number, ph['score'], self.subject(ph['subject']))
                ).bold = True
                ans.append(self._correct(ph['choices']))
                text = '正确答案：%s\n' % (self._correct(ph['choices']))
                p1.add_run(text)

                explanation = ph['explanation'].replace('<br>', '\n')
                if '<img' in explanation:
                    urls = re.findall('<img src="(.*?)".*?>', explanation)
                    texts = re.split('<img.*?>', explanation)
                    p1.add_run(texts[0])
                    for t, u in zip(texts[1:], urls):
                        r = p1.add_run()
                        inline_shape = r.add_picture(self.download(u))
                        inline_shape.height = int(inline_shape.height * 0.4)
                        inline_shape.width = int(inline_shape.width * 0.5)
                        p1.add_run(t)
                        time.sleep(1)
                else:
                    p1.add_run('%s\n' % explanation)
                p1.style = 'Normal'
                problem_number += 1
            time.sleep(2)

        i = 1
        p1 = document.add_paragraph()
        for index in range(0, len(ans), 5):
            p1.add_run('%s-%s %s    ' % (i, i + 4, ''.join(ans[index:index + 5]))).bold = True
            i += 5

        if name:
            document.save(name + '.docx')
        else:
            document.save(str(uuid.uuid4().hex) + '.docx')
        return True


def read_config():
    with open("./config.json", 'r', encoding='UTF-8') as fp:
        return json.load(fp)


if __name__ == '__main__':
    config = read_config()
    client = init_client(config['auth'])
    process = Process(client, config['user_id'])
    paper = Paper(client, process.process)
    paper.run(config['data'])
